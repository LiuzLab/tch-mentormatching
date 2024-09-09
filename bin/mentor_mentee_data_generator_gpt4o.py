import openai
import pandas as pd
from openai import OpenAI
import os
import random
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from langchain.document_loaders import PyPDFLoader
from langchain.chains.question_answering import load_qa_chain
from langchain_openai import ChatOpenAI
import asyncio
from docx import Document
from langchain.docstore.document import Document as LangchainDocument
from langchain_core.pydantic_v1 import BaseModel, Field
from typing import cast, List, Literal


class Mentee(BaseModel):
    thoughts: List[str] = Field(description="Reasoning Steps to Extract the information")
    is_assistant_professor: bool = Field(description="True if the mentee is assistant professor or above.")
    education: Literal["BS and Master and PhD's Degree", "BS and Master's Degree", "Only BS Degree"] = Field()

    # NOTE(JL): the below is poorly extracted.
    # n_papers: int = Field(description="The number of publications listed in the CV")
    # n_experiences: int = Field(description="The number of experiences listed in the CV")
    # n_skills: int = Field(description="The number of skills listed in the CV")
    # n_volunteering: int = Field(description="The number of volunteering activities listed in the CV")

load_dotenv()

client = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.5,
    max_tokens=3000,
    api_key=os.getenv("OPENAI_API_KEY"),
)

async def generate_samples(prompt, mentor_profile_documents):
    chain = load_qa_chain(client, verbose=True)
    response = await chain.arun(input_documents=mentor_profile_documents, question=prompt)
    return response

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return ' '.join(full_text)


def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for i in range(min(2, len(reader.pages))):
        page = reader.pages[i]
        text += page.extract_text()
    text = " ".join(text.split())
    return text


def load_pdf(file_path):
    loader = PyPDFLoader(file_path)
    return loader.load_and_split()


def load_docx(file_path):
    text = extract_text_from_docx(file_path)
    return [LangchainDocument(page_content=text, metadata={"source": file_path})]


async def generate_mock_cv(file_path):
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.docx':
        file_text = extract_text_from_docx(file_path)
        mentor_profile_documents = load_docx(file_path)
    elif file_extension == '.pdf':
        file_text = extract_text_from_pdf(file_path)
        mentor_profile_documents = load_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}. Please use .docx or .pdf")

    structured_llm = client.with_structured_output(Mentee, strict=True)
    mentee: Mentee = cast(Mentee, structured_llm.invoke(f"""Extract the mentee's information from the CV. CV: {file_text}"""))
    print(mentee)
    return mentee, file_text


if __name__ == "__main__":
    # Add any test code here if needed
    pass
